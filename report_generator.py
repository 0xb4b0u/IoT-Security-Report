# generate_report.py
import os
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from classes import CVE


def generate_report(cve: CVE) -> None:
    """
    Génère et sauvegarde un rapport au format Word pour une vulnérabilité (CVE) donnée.

    Le rapport inclut :
    - Les informations générales de la CVE (identifiant, date de publication, score CVSS, etc.)
    - Un résumé exécutif avec la description de la vulnérabilité
    - Un tableau récapitulatif des machines impactées

    :param cve: Instance de CVE contenant les informations à intégrer dans le rapport.
    """
    report = Document()

    # Titre du rapport et identifiant de la CVE
    title = report.add_heading("Rapport Cyber-résilience", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cve_heading = report.add_heading(f"{cve.cve_id}", level=0)
    cve_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Informations générales : Date de publication et score CVSS
    publication_date_paragraph = report.add_paragraph()
    publication_date_paragraph.add_run("Date de publication: ").bold = True
    publication_date_paragraph.add_run(f"{cve.publication_date}")

    cvss_severity = cve.cvss_info.get("sévérité", "N/A")
    cvss_score_paragraph = report.add_paragraph()
    cvss_score_paragraph.add_run("Score CVSS: ").bold = True
    cvss_score_paragraph.add_run(f"{cve.cvss_score} ({cvss_severity})")

    # Affichage des autres informations CVSS s'il y en a
    if cve.cvss_info:
        for key, value in cve.cvss_info.items():
            formatted_key = key.replace("_", " ").capitalize()
            cvss_info_paragraph = report.add_paragraph()
            cvss_info_paragraph.add_run(f"{formatted_key}: ").bold = True
            cvss_info_paragraph.add_run(f"{value}".lower())

    # Résumé exécutif
    report.add_heading("Résumé exécutif", level=1)

    # Description de la CVE avec une numérotation personnalisée
    description_paragraph = report.add_paragraph(
        "Description de la CVE: ", style="List Number 2"
    )
    description_paragraph.add_run("\n")
    description_paragraph.add_run(f"{cve.description}")

    # Section d'impact TODO: à compléter
    impact_paragraph = report.add_paragraph("Impact de la CVE: ", style="List Number 2")
    impact_paragraph.add_run("\n")
    impact_paragraph.add_run("TODO")

    # Tableau des machines impactées
    report.add_heading("Tableau des machines impactées par la CVE", level=1)
    report.add_paragraph(f"{len(cve.devices)} machine(s) impactée(s)")

    if cve.devices:
        # Création d'un tableau avec 5 colonnes
        table = report.add_table(rows=1, cols=5)
        table.style = "Table Grid"

        header_cells = table.rows[0].cells
        header_cells[0].text = "Nom"
        header_cells[1].text = "IP"
        header_cells[2].text = "Profil"
        header_cells[3].text = "Site"
        header_cells[4].text = "Dernière détection"

        # Ajout d'une ligne pour chaque machine
        for device in cve.devices:
            row_cells = table.add_row().cells
            row_cells[0].text = device.name
            row_cells[1].text = device.ip
            row_cells[2].text = device.profile
            row_cells[3].text = device.site
            row_cells[4].text = str(device.last_detected_date)
    else:
        report.add_paragraph("Aucune machine détectée.")

    if cve.false_positive:
        # Tableau des machines "faux positifs"
        report.add_heading("Tableau faux positif sous conditions", level=1)
        report.add_paragraph(f"{len(cve.false_positive)} machine(s) non détectée(s) sur le réseau depuis + d'un mois")

        # Création d'un tableau avec 5 colonnes
        table = report.add_table(rows=1, cols=5)
        table.style = "Table Grid"

        header_cells = table.rows[0].cells
        header_cells[0].text = "Nom"
        header_cells[1].text = "IP"
        header_cells[2].text = "Profil"
        header_cells[3].text = "Site"
        header_cells[4].text = "Dernière détection"

        # Ajout d'une ligne pour chaque machine
        for device in cve.false_positive:
            row_cells = table.add_row().cells
            row_cells[0].text = device.name
            row_cells[1].text = device.ip
            row_cells[2].text = device.profile
            row_cells[3].text = device.site
            row_cells[4].text = str(device.last_detected_date)
    else:
        report.add_paragraph("Aucune machine détectée.")

    # Création du répertoire 'reports' si besoin et sauvegarde du document
    reports_dir = "reports"
    os.makedirs(reports_dir, exist_ok=True)
    docx_path = os.path.join(reports_dir, f"{cve.cve_id}_report.docx")
    report.save(docx_path)
